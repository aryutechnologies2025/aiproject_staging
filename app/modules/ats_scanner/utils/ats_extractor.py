from __future__ import annotations

import os
import io
import re
import logging
from typing import List, Dict, Optional

from fastapi import UploadFile
from llama_cloud import AsyncLlamaCloud

logger = logging.getLogger(__name__)

LLAMA_CLOUD_API_KEY = os.getenv("LLAMA_CLOUD_API_KEY")

_BULLET_CHARS = frozenset("•·▪▫►▶→▷➤➢➣✦✧✓✔✗✘❑❒❖⁃‣◦")

_SECTION_HEADING_KEYWORDS = {
    "experience", "professional experience", "work experience",
    "employment", "employment history", "work history", "career history",
    "career", "internship", "internships", "industrial training",
    "training", "professional background", "positions held",
    "relevant experience", "prior experience", "work background",
    "career background", "occupation", "apprenticeship", "placement",
    "scholastic background", "scholastic details", "scholastic achievements",
    "education", "educational background", "academic background",
    "qualifications", "academic qualifications", "academic",
    "skills", "technical skills", "tech stack", "technical stack",
    "core competencies", "competencies", "expertise", "technologies",
    "tools", "tools & technologies", "tools and technologies",
    "programming languages", "software skills", "proficiencies",
    "projects", "key projects", "personal projects", "project experience",
    "portfolio",
    "summary", "professional summary", "objective", "career objective",
    "profile", "about me", "overview", "highlights",
    "contact", "contact information", "contact details",
    "certifications", "certificates", "licenses", "credentials", "courses",
    "languages", "language proficiency",
    "awards", "achievements", "honors", "honours",
    "volunteer", "volunteering", "publications", "research",
    "hobbies", "interests", "references",
}

_ALLCAPS_RE = re.compile(r"^[A-Z][A-Z\s&/\-]{2,53}[A-Z]$")
_BULLET_RE  = re.compile(r"^[\s]*[•\-\*\u2022\u2023►▶→✓✔]+\s*")


def _sanitise(text: str) -> str:
    if not text:
        return ""
    return text.encode("utf-8", errors="replace").decode("utf-8", errors="replace")


def _normalise_bullets(text: str) -> str:
    result: List[str] = []
    for line in text.split("\n"):
        stripped = line.lstrip()
        if stripped and stripped[0] in _BULLET_CHARS:
            indent  = line[: len(line) - len(stripped)]
            content = stripped[1:].lstrip()
            result.append(f"{indent}- {content}")
        else:
            result.append(line)
    return "\n".join(result)


def _repair_broken_table(text: str) -> str:
    pipe_line = re.compile(r"^\s*\|.*\|\s*$")
    align_row = re.compile(r"^\s*\|[\s:\-|]+\|\s*$")
    lines     = text.split("\n")
    result:     List[str] = []
    table_buf:  List[str] = []
    headers:    List[str] = []

    def flush_table():
        nonlocal headers
        if not table_buf:
            return
        header_row = next((r for r in table_buf if not align_row.match(r)), None)
        if header_row:
            headers = [c.strip() for c in header_row.strip("|").split("|") if c.strip()]
        for row in table_buf:
            if align_row.match(row) or row == header_row:
                continue
            cells = [c.strip() for c in row.strip("|").split("|")]
            if len(cells) == len(headers) and headers:
                for h, c in zip(headers, cells):
                    if c:
                        result.append(f"- **{h}**: {c}")
            else:
                for c in cells:
                    if c:
                        result.append(f"- {c}")
        table_buf.clear()
        headers.clear()

    for line in lines:
        if pipe_line.match(line) or (line.strip().startswith("|") and "|" in line[1:]):
            table_buf.append(line)
        else:
            if table_buf:
                flush_table()
            result.append(line)
    flush_table()
    return "\n".join(result)


def _merge_fragments(text: str) -> str:
    text = re.sub(r"-\s*\n\s*(\w)", r"\1", text)
    text = re.sub(r"(https?://[^\s]+)\s*\n\s*([^\s]{2,40})", lambda m: m.group(1) + m.group(2), text)
    lines  = text.split("\n")
    merged: List[str] = []
    for line in lines:
        s = line.strip()
        if (merged and s and len(s) <= 2
                and not re.match(r"^[-•*#|]", s)
                and not s.isdigit()):
            merged[-1] = merged[-1].rstrip() + " " + s
        else:
            merged.append(line)
    return "\n".join(merged)


async def extract_resume_markdown(file: UploadFile) -> str:
    filename = file.filename or "resume.pdf"
    await file.seek(0)
    file_bytes = await file.read()
    await file.seek(0)

    content_type = file.content_type or (
        "application/pdf" if filename.lower().endswith(".pdf")
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    if LLAMA_CLOUD_API_KEY:
        try:
            md = await _llamaparse_to_markdown(file_bytes, filename, content_type)
            if md and len(md.strip()) > 100:
                logger.info(f"LlamaParse markdown: {len(md)} chars")
                return _sanitise(_post_process_markdown(md))
            logger.warning("LlamaParse short/empty, using local fallback")
        except Exception as e:
            logger.warning(f"LlamaParse failed: {e}, using local fallback")

    md = _local_extract_markdown(file_bytes, filename)
    logger.info(f"Local fallback markdown: {len(md)} chars")
    return _sanitise(_post_process_markdown(md))


def _post_process_markdown(md: str) -> str:
    md = _normalise_bullets(md)
    md = _merge_fragments(md)
    md = _repair_broken_table(md)
    md = re.sub(r"\n{3,}", "\n\n", md)
    md = "\n".join(l.rstrip() for l in md.split("\n"))
    return md.strip()


async def _llamaparse_to_markdown(file_bytes: bytes, filename: str, content_type: str) -> str:
    client   = AsyncLlamaCloud(api_key=LLAMA_CLOUD_API_KEY)
    uploaded = await client.files.create(
        file=(filename, file_bytes, content_type), purpose="parse"
    )
    result = await client.parsing.parse(
        file_id=uploaded.id, tier="agentic", version="latest", expand=["items"]
    )
    items = _flat_items(result)
    return _items_to_markdown(items) if items else ""


def _flat_items(result) -> List:
    items_obj = getattr(result, "items", None)
    if not items_obj:
        return []
    pages = getattr(items_obj, "pages", None) or []
    flat  = []
    for page in pages:
        pnum = getattr(page, "page_number", 1)
        for item in (getattr(page, "items", None) or []):
            if not hasattr(item, "page_number"):
                setattr(item, "page_number", pnum)
            flat.append(item)
    return flat


def _items_to_markdown(items: List) -> str:
    lines:     List[str] = []
    prev_page: Optional[int] = None

    for item in items:
        pnum = getattr(item, "page_number", 1)
        if prev_page is not None and pnum != prev_page:
            lines.append("")
        prev_page = pnum

        text = (
            str(item.md).strip()    if getattr(item, "md",    None) else
            str(item.value).strip() if getattr(item, "value", None) else
            str(item.text).strip()  if getattr(item, "text",  None) else ""
        )
        if not text:
            for n in _nested_texts(item):
                lines.append(f"- {n}")
            continue

        btype = str(getattr(item, "type", "text") or "text").lower()

        if btype in ("heading", "h1", "h2", "h3", "h4", "title"):
            clean_text = re.sub(r'^#+\s*', '', text).strip()
            lines.append(f"\n## {clean_text}\n")
        elif btype in ("list", "bullet", "li"):
            clean_text = re.sub(r'^[-•*]\s*', '', text).strip()
            lines.append(f"- {clean_text}")
            for n in _nested_texts(item):
                lines.append(f"- {n}")
        elif btype == "table":
            lines.append(_repair_broken_table(text))
        else:
            lines.append(text)

    return "\n".join(lines)


def _nested_texts(item) -> List[str]:
    out = []
    for sub in (getattr(item, "items", None) or []):
        v = (
            str(sub.value).strip() if getattr(sub, "value", None) else
            str(sub.md).strip()    if getattr(sub, "md",    None) else
            str(sub.text).strip()  if getattr(sub, "text",  None) else ""
        )
        if v:
            out.append(v)
    return out


def _local_extract_markdown(file_bytes: bytes, filename: str) -> str:
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return _pdf_to_markdown(file_bytes)
    elif fname.endswith(".docx"):
        return _docx_to_markdown(file_bytes)
    return ""


def _pdf_to_markdown(file_bytes: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber not installed")
        return ""

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_md: List[str] = []
            for page in pdf.pages:
                pages_md.append(_page_to_markdown(page))
            return "\n\n".join(p for p in pages_md if p.strip())
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


def _page_to_markdown(page) -> str:
    words = page.extract_words(x_tolerance=3, y_tolerance=3,
                               keep_blank_chars=False, use_text_flow=False)
    if not words:
        return ""

    split = _density_column_split(words, page.width)

    if split is None:
        lines = _words_to_lines(words)
        return _lines_to_markdown(lines)
    else:
        left_words  = [w for w in words if w["x0"] < split]
        right_words = [w for w in words if w["x0"] >= split]
        left_md     = _lines_to_markdown(_words_to_lines(left_words))
        right_md    = _lines_to_markdown(_words_to_lines(right_words))
        return left_md + "\n\n" + right_md


def _density_column_split(words: List[Dict], page_width: float) -> Optional[float]:
    lo = page_width * 0.25
    hi = page_width * 0.75

    bucket: Dict[int, int] = {}
    for w in words:
        x0 = w["x0"]
        if lo <= x0 <= hi:
            b = int(x0 / 5) * 5
            bucket[b] = bucket.get(b, 0) + 1

    if not bucket:
        return None

    xs = sorted(bucket)
    best_gap_mid   = None
    best_gap_score = 9999

    for x in xs:
        window_count = sum(
            bucket.get(b, 0)
            for b in range(int(x / 5) * 5 - 2, int(x / 5) * 5 + 6)
        )
        if window_count < best_gap_score:
            best_gap_score = window_count
            best_gap_mid   = x

    if best_gap_mid is None:
        return None

    left_dense  = sum(bucket.get(b, 0) for b in range(0,               int(best_gap_mid / 5) * 5) if lo / 5 < b)
    right_dense = sum(bucket.get(b, 0) for b in range(int(best_gap_mid / 5) * 5, 9999)             if b < hi / 5 * 5)

    if left_dense < 5 or right_dense < 5:
        return None

    avg_side_density = (left_dense + right_dense) / (len(xs) or 1)
    if best_gap_score > avg_side_density * 0.3:
        return None

    return float(best_gap_mid + 5)


def _words_to_lines(words: List[Dict], ytol: float = 4) -> List[str]:
    if not words:
        return []
    sw  = sorted(words, key=lambda w: (round(w["top"] / ytol) * ytol, w["x0"]))
    lines: List[str] = []
    cur:   List[Dict] = [sw[0]]
    cy                = sw[0]["top"]

    for w in sw[1:]:
        if abs(w["top"] - cy) <= ytol:
            cur.append(w)
        else:
            lines.append(" ".join(x["text"] for x in sorted(cur, key=lambda x: x["x0"])))
            cur, cy = [w], w["top"]

    if cur:
        lines.append(" ".join(x["text"] for x in sorted(cur, key=lambda x: x["x0"])))

    return lines


def _lines_to_markdown(lines: List[str]) -> str:
    md: List[str] = []

    for line in lines:
        s = line.strip()
        if not s:
            md.append("")
            continue

        s_lower = s.lower().rstrip(":").strip()
        if s_lower in _SECTION_HEADING_KEYWORDS:
            md.append(f"\n## {s.rstrip(':')}\n")
            continue

        if _ALLCAPS_RE.match(s) and len(s) <= 55:
            if not re.search(r"\d", s[:4]):
                md.append(f"\n## {s.title()}\n")
                continue

        if len(s.split()) <= 6:
            for kw in _SECTION_HEADING_KEYWORDS:
                if kw in s_lower and len(kw) >= 6:
                    md.append(f"\n## {s.rstrip(':')}\n")
                    break
            else:
                if _BULLET_RE.match(s):
                    clean = _BULLET_RE.sub("", s).strip()
                    md.append(f"- {clean}")
                else:
                    md.append(s)
            continue

        if _BULLET_RE.match(s):
            clean = _BULLET_RE.sub("", s).strip()
            md.append(f"- {clean}")
            continue

        md.append(s)

    return "\n".join(md)


def _docx_to_markdown(file_bytes: bytes) -> str:
    try:
        import docx as python_docx
        doc   = python_docx.Document(io.BytesIO(file_bytes))
        lines: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    lines.append(row_text)
        return _lines_to_markdown(lines)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""