# /home/aryu_user/Arun/aiproject_staging/app/modules/ats_scanner/utils/text_extraction.py
"""
Production-Grade Text Extraction v2
Enhanced PDF/DOCX parsing with better section detection
Specifically improved education extraction + summary fallback fix

v2.1 — Added UTF-8 sanitisation to handle resumes from Windows/Word
       that embed Latin-1 or CP-1252 encoded characters (ö, ü, é etc.)
"""

import pdfplumber
import docx
from fastapi import UploadFile
import re
import logging
from typing import Dict, List, Tuple, Optional

logger = logging.getLogger(__name__)

# =====================================================
# SECTION DETECTION PATTERNS
# =====================================================

SECTION_PATTERNS = {
    "education": [
        r"(?:^|\n)\s*(?:EDUCATION|ACADEMIC|QUALIFICATION|DEGREE|SCHOOL|UNIVERSITY)\s*(?:$|\n)",
        r"(?:^|\n)\s*(?:B\.?S|B\.?A|M\.?S|M\.?A|PhD|M\.?B\.?A|Bachelor|Master|Associate|Diploma)\s+(?:in|of)",
    ],
    "experience": [
        r"(?:^|\n)\s*(?:EXPERIENCE|WORK EXPERIENCE|PROFESSIONAL EXPERIENCE|EMPLOYMENT|CAREER)\s*(?:$|\n)",
        r"(?:^|\n)\s*(?:Senior|Junior|Lead|Principal|Manager|Engineer|Developer|Analyst|Specialist)\s+",
    ],
    "skills": [
        r"(?:^|\n)\s*(?:SKILLS|TECHNICAL SKILLS|CORE COMPETENCIES|EXPERTISE|TECHNOLOGIES)\s*(?:$|\n)",
        r"(?:^|\n)\s*(?:Languages?|Frameworks?|Tools?|Platforms?|Databases?|Software)\s*:",
    ],
    "summary": [
        r"(?:^|\n)\s*(?:SUMMARY|OBJECTIVE|PROFESSIONAL SUMMARY|PROFILE|ABOUT ME|OVERVIEW|INTRODUCTION)\s*(?:$|\n)",
    ]
}

# Patterns that indicate a line is NOT summary content (contact info, section headers, etc.)
SKIP_LINE_PATTERNS = [
    r"@",                          # Email
    r"linkedin\.com",
    r"github\.com",
    r"\+?\d{10,}",                 # Phone numbers
    r"\+91|\+1|\+44",              # Country codes
    r"^\s*http",                   # URLs
    r"^\s*www\.",
    r"^\s*[A-Z\s]{2,30}\s*$",     # All-caps short lines (section headers)
]

# Education degree patterns
EDUCATION_PATTERNS = {
    "degree": [
        r"(?:B\.?S\.?|B\.?A\.?|Bachelor|M\.?S\.?|M\.?A\.?|Master|Ph\.?D\.?|Doctor|M\.?B\.?A\.?|MBA|Associate)",
        r"(?:in|of)\s+([A-Z][a-zA-Z\s&]+?)(?:,|;|$|\n)",
    ],
    "institution": [
        r"(?:University|College|Institute|School|Academy|Technical)",
    ],
    "year": [
        r"(?:20\d{2}|19\d{2})",
    ]
}


# =====================================================
# UTF-8 SANITISATION HELPER
# =====================================================

def _sanitise(text: str) -> str:
    """
    Sanitise extracted text to pure UTF-8.

    PDFs and DOCX files created on Windows often embed characters in
    CP-1252 / Latin-1 encoding (e.g. ö = 0xf6, ü = 0xfc, é = 0xe9).
    When pdfplumber or python-docx reads them they may come through as
    raw byte values that later crash FastAPI's JSON encoder.

    Using errors="replace" turns any un-encodable byte into the Unicode
    replacement character (U+FFFD) rather than raising an exception.
    """
    if not text:
        return text
    return text.encode("utf-8", errors="replace").decode("utf-8", errors="replace")


# =====================================================
# TEXT EXTRACTION ENGINE
# =====================================================

class TextExtractionEngine:
    """Production-grade text extraction with section detection"""

    def __init__(self):
        self.raw_text = ""
        self.sections = {}

    async def extract_all(self, file: UploadFile) -> Dict:
        """
        Extract all text and structure from file

        Returns:
            {
                "raw_text": "Full extracted text",
                "sections": {
                    "education": "Education section text",
                    "experience": "Experience section text",
                    ...
                },
                "metadata": {
                    "file_type": "pdf",
                    "pages": 2,
                    ...
                }
            }
        """

        filename = file.filename.lower() if file.filename else ""

        try:
            if filename.endswith(".pdf"):
                result = await self._extract_pdf_structured(file)
            elif filename.endswith(".docx"):
                result = await self._extract_docx_structured(file)
            elif filename.endswith(".doc"):
                result = await self._extract_docx_structured(file)
            else:
                raise ValueError(f"Unsupported file type: {filename}")

            logger.info(f"Text extraction successful. Found {len(result['sections'])} sections")
            return result

        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            raise

    async def extract_text(self, file: UploadFile) -> str:
        """
        Legacy method - extract raw text only
        For backward compatibility
        """
        result = await self.extract_all(file)
        return result["raw_text"]

    # =========== PDF EXTRACTION ===========

    async def _extract_pdf_structured(self, file: UploadFile) -> Dict:
        """Extract from PDF with structure detection"""

        text_pages = []
        metadata = {
            "file_type": "pdf",
            "pages": 0,
            "has_tables": False,
            "has_images": False
        }

        try:
            # Reset file pointer before reading
            await file.seek(0)

            with pdfplumber.open(file.file) as pdf:
                metadata["pages"] = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text() or ""
                    tables = page.extract_tables()

                    # Detect tables
                    if tables:
                        metadata["has_tables"] = True
                        table_text = self._extract_pdf_tables(tables)
                        page_text += "\n" + table_text

                    # Detect images
                    if page.images:
                        metadata["has_images"] = True

                    if page_text.strip():
                        text_pages.append(page_text)

            raw_text = "\n".join(text_pages)

        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise

        # ── FIX: sanitise before any further processing ──────────────────────
        # PDFs from Windows / Word may embed Latin-1 or CP-1252 bytes that
        # crash FastAPI's JSON encoder downstream.
        raw_text = _sanitise(raw_text)

        # Detect sections
        sections = self._detect_sections(raw_text)

        return {
            "raw_text": raw_text,
            "sections": sections,
            "metadata": metadata
        }

    def _extract_pdf_tables(self, tables) -> str:
        """Convert PDF tables to text"""
        table_text = []

        for table in tables:
            for row in table:
                row_str = " | ".join([str(cell) if cell else "" for cell in row])
                if row_str.strip():
                    table_text.append(row_str)

        return "\n".join(table_text)

    # =========== DOCX EXTRACTION ===========

    async def _extract_docx_structured(self, file: UploadFile) -> Dict:
        """Extract from DOCX with structure detection"""

        text_parts = []
        metadata = {
            "file_type": "docx",
            "paragraphs": 0,
            "has_tables": False,
            "has_images": False,
            "has_formatting": False
        }

        try:
            await file.seek(0)
            doc = docx.Document(file.file)

            metadata["paragraphs"] = len(doc.paragraphs)
            metadata["has_tables"] = len(doc.tables) > 0

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()

                if text:
                    # Check for formatting
                    for run in para.runs:
                        if run.bold or run.italic or run.underline:
                            metadata["has_formatting"] = True
                            break

                    text_parts.append(text)

            # Extract tables
            if doc.tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text_parts.append(row_text)

            # Check for images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    metadata["has_images"] = True

            raw_text = "\n".join(text_parts)

        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise

        # ── FIX: sanitise before any further processing ──────────────────────
        # DOCX files saved on Windows may contain CP-1252 characters.
        raw_text = _sanitise(raw_text)

        # Detect sections
        sections = self._detect_sections(raw_text)

        return {
            "raw_text": raw_text,
            "sections": sections,
            "metadata": metadata
        }

    # =========== SECTION DETECTION ===========

    def _detect_sections(self, text: str) -> Dict[str, str]:
        """
        Intelligently detect resume sections.
        Falls back to heuristic detection for unlabeled summaries.
        """

        sections = {}

        for section_name, patterns in SECTION_PATTERNS.items():
            section_text = self._extract_section(text, section_name, patterns)

            if section_text and section_text.strip():
                sections[section_name] = section_text.strip()

        # =====================================================
        # FIX: Fallback summary detection
        # If no labeled SUMMARY / PROFILE / OBJECTIVE section found,
        # attempt to extract the first substantive paragraph that
        # appears before any known section headers.
        # =====================================================
        if "summary" not in sections or not sections.get("summary", "").strip():
            fallback_summary = self._extract_fallback_summary(text)
            if fallback_summary:
                sections["summary"] = fallback_summary
                logger.info("Summary extracted via fallback heuristic (no explicit header found)")

        return sections

    def _extract_fallback_summary(self, text: str) -> str:
        """
        Heuristic: Scan the first ~30 lines of the resume.
        Skip contact info, headers, and blank lines.
        Collect the first block of prose text (5+ words) as summary.
        Stop when a known section header is encountered.
        """

        lines = text.split("\n")
        summary_lines = []
        found_start = False
        section_header_pattern = re.compile(
            r"^\s*(?:EDUCATION|EXPERIENCE|WORK EXPERIENCE|SKILLS|TECHNICAL SKILLS|"
            r"PROFESSIONAL EXPERIENCE|EMPLOYMENT|PROJECTS|CERTIFICATIONS|"
            r"ACHIEVEMENTS|AWARDS|INTERESTS|REFERENCES|CAREER)\s*$",
            re.IGNORECASE
        )

        for line in lines[:40]:  # Only scan first 40 lines
            stripped = line.strip()

            # Stop if we hit a known section header
            if section_header_pattern.match(stripped):
                break

            # Skip empty lines between blocks (but don't break)
            if not stripped:
                if found_start and summary_lines:
                    # One blank line ends the summary block
                    break
                continue

            # Skip lines matching contact info / skip patterns
            is_skip = False
            for pattern in SKIP_LINE_PATTERNS:
                if re.search(pattern, stripped, re.IGNORECASE):
                    is_skip = True
                    break

            if is_skip:
                continue

            # Skip pure all-caps short lines (likely name or section headers)
            if stripped.isupper() and len(stripped) < 50:
                continue

            # Skip name-like lines at very top (single line, Title Case, no verbs)
            if not found_start and len(stripped.split()) <= 4 and stripped.istitle():
                continue

            # Accept substantive lines
            if len(stripped.split()) >= 5:
                found_start = True
                summary_lines.append(stripped)

            # Collect up to 3 lines for summary
            if len(summary_lines) >= 3:
                break

        return " ".join(summary_lines).strip()

    def _extract_section(self, text: str, section_name: str, patterns: List[str]) -> str:
        """Extract section content using patterns"""

        # Find section start
        start_pos = -1

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                start_pos = match.start()
                break

        if start_pos == -1:
            return ""

        # Find section end (next section header or end of text)
        end_pos = len(text)

        # Look for next section header
        remaining_text = text[start_pos + 50:]  # Skip current header

        for other_section, other_patterns in SECTION_PATTERNS.items():
            if other_section == section_name:
                continue

            for pattern in other_patterns:
                match = re.search(pattern, remaining_text, re.IGNORECASE | re.MULTILINE)
                if match:
                    potential_end = start_pos + 50 + match.start()
                    if potential_end < end_pos:
                        end_pos = potential_end

        section_text = text[start_pos:end_pos]

        # Remove the header line itself
        section_text = re.sub(
            r"^[^\n]*(?:SUMMARY|OBJECTIVE|PROFESSIONAL SUMMARY|PROFILE|ABOUT ME|OVERVIEW|"
            r"EDUCATION|ACADEMIC|EXPERIENCE|WORK EXPERIENCE|SKILLS|TECHNICAL SKILLS|"
            r"EXPERTISE|TECHNOLOGIES)[^\n]*\n?",
            "",
            section_text,
            count=1,
            flags=re.IGNORECASE
        )

        return section_text.strip()

    # =========== EDUCATION EXTRACTION ===========

    def extract_education_entries(self, text: str) -> List[Dict]:
        """
        Extract individual education entries from text

        Returns:
            [
                {
                    "degree": "B.S. Computer Science",
                    "institution": "MIT",
                    "year": "2020",
                    "gpa": "3.8",
                    "raw_text": "Full entry text"
                },
                ...
            ]
        """

        entries = []

        if not text or not text.strip():
            return entries

        # Split into lines and group into entries
        lines = text.split("\n")
        current_entry = []

        for line in lines:
            line = line.strip()

            if not line:
                if current_entry:
                    entry_text = " ".join(current_entry)
                    parsed = self._parse_education_entry(entry_text)
                    if parsed:
                        entries.append(parsed)
                    current_entry = []
                continue

            # Check if this line looks like the start of a new education entry
            is_new_entry = re.match(
                r"^(B\.?S|B\.?A|B\.?Tech|B\.?E|M\.?S|M\.?A|M\.?Tech|Ph\.?D|MBA|Bachelor|Master|Associate|Diploma|Certificate)",
                line,
                re.IGNORECASE
            )

            if is_new_entry and current_entry:
                entry_text = " ".join(current_entry)
                parsed = self._parse_education_entry(entry_text)
                if parsed:
                    entries.append(parsed)
                current_entry = []

            current_entry.append(line)

        # Flush the last entry
        if current_entry:
            entry_text = " ".join(current_entry)
            parsed = self._parse_education_entry(entry_text)
            if parsed:
                entries.append(parsed)

        logger.info(f"Extracted {len(entries)} education entries")
        return entries

    def _parse_education_entry(self, text: str) -> Optional[Dict]:
        """Parse individual education entry into structured dict"""

        if not text or len(text.strip()) < 5:
            return None

        entry = {
            "degree": "",
            "institution": "",
            "year": "",
            "gpa": "",
            "raw_text": text
        }

        # Extract degree
        degree_match = re.search(
            r"(B\.?S\.?|B\.?A\.?|B\.?Tech\.?|B\.?E\.?|Bachelor[^,\n]*|"
            r"M\.?S\.?|M\.?A\.?|M\.?Tech\.?|Master[^,\n]*|"
            r"Ph\.?D\.?[^,\n]*|Doctor[^,\n]*|"
            r"M\.?B\.?A\.?|MBA|Associate[^,\n]*|"
            r"Diploma[^,\n]*|Certificate[^,\n]*|Bootcamp[^,\n]*)",
            text,
            re.IGNORECASE
        )
        if degree_match:
            entry["degree"] = degree_match.group(0).strip().rstrip(",;")

        # Extract institution
        institution_match = re.search(
            r"(?:from\s+|at\s+|,\s*|\|\s*)([\w\s&'\-\.]+?(?:University|College|Institute|School|Academy|Technical|Tech)[\w\s&'\-\.]*?)(?:\s*,|\s*\||\s*\d{4}|$)",
            text,
            re.IGNORECASE
        )
        if institution_match:
            entry["institution"] = institution_match.group(1).strip()
        else:
            # Fallback: look for known institution keywords anywhere
            inst_match = re.search(
                r"([\w\s&'\-\.]+(?:University|College|Institute|School|Academy)[\w\s&'\-\.]*)",
                text,
                re.IGNORECASE
            )
            if inst_match:
                entry["institution"] = inst_match.group(1).strip()

        # Extract graduation year
        year_match = re.search(r"(20\d{2}|19\d{2})", text)
        if year_match:
            entry["year"] = year_match.group(0)

        # Extract GPA
        gpa_match = re.search(r"GPA[:\s]*([0-4]\.[0-9]{1,2})", text, re.IGNORECASE)
        if gpa_match:
            entry["gpa"] = gpa_match.group(1)

        # Only return if we extracted at least degree or institution
        if entry["degree"] or entry["institution"]:
            return entry

        return None


# =====================================================
# LEGACY ASYNC FUNCTIONS (backward compatibility)
# =====================================================

async def extract_text(file: UploadFile) -> str:
    """
    Extract text from PDF/DOCX — backward compatible wrapper.
    Output is always sanitised UTF-8 (no Latin-1 / CP-1252 bytes).
    """
    engine = TextExtractionEngine()
    result = await engine.extract_all(file)
    return result["raw_text"]   # already sanitised inside extract_all


async def extract_pdf_text(file: UploadFile) -> str:
    """Extract text from PDF"""
    engine = TextExtractionEngine()
    result = await engine._extract_pdf_structured(file)
    return result["raw_text"]


async def extract_docx_text(file: UploadFile) -> str:
    """Extract text from DOCX"""
    engine = TextExtractionEngine()
    result = await engine._extract_docx_structured(file)
    return result["raw_text"]