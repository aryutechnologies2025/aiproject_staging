"""
parser_service.py — Unified resume parsing service for resume_builder.

Handles:
- Universal multi-format ingestion (PDF, DOCX, DOC, TXT) directly from raw bytes.
- Direct Gemini document understanding for PDF with PyMuPDF layout fallback.
- Python-docx paragraph and table extraction for Word documents.
- Flexible extractor_output compatibility (raw_items, text dictionaries, raw strings).
- Ingress request validation, sanitization, and structured audit logging.
"""

import io
import logging
from typing import Any, Dict, Optional
from fastapi import HTTPException, Request

from app.modules.resume_builder.security_manager import SecurityManager
from app.modules.resume_builder.input_sanitizer import InputSanitizer
from app.modules.resume_builder.request_validator import RequestValidator
from app.modules.resume_builder.security_audit_logger import SecurityAuditLogger
from app.modules.resume_builder.ai_parser import ImprovedUniversalResumeParser

logger = logging.getLogger("resume_builder.parser_service")


async def parse_resume_with_ai(
    extractor_output: Optional[Any] = None,
    file_bytes: Optional[bytes] = None,
    filename: Optional[str] = None,
    content_type: Optional[str] = None,
    request: Optional[Request] = None,
    user_id: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Main entry point for parsing resumes with Gemini and Deterministic AST fallback.
    Accepts raw file_bytes across all supported document formats (PDF, DOCX, DOC, TXT)
    or extractor_output payloads.
    """
    try:
        client_ip = "unknown"
        resolved_user_id = user_id or "anonymous"
        safe_filename = filename or "resume.pdf"

        if request:
            resolved_user_id = request.headers.get("X-User-ID", resolved_user_id)
            client_ip = RequestValidator.get_client_ip(request)

            # Validate request
            is_valid, error = RequestValidator.validate_request(request)
            if not is_valid:
                logger.warning(f"[ParserService] Request validation failed: {error}")
                SecurityAuditLogger.log_security_violation(
                    violation_type="invalid_request",
                    client_ip=client_ip,
                    user_id=resolved_user_id,
                    details=error or "Validation failed",
                )
                raise HTTPException(status_code=400, detail=error)

            # Rate limit check
            is_allowed, error = SecurityManager.check_rate_limit(resolved_user_id, client_ip)
            if not is_allowed:
                logger.warning(f"[ParserService] Rate limit exceeded for {client_ip}: {error}")
                SecurityAuditLogger.log_security_violation(
                    violation_type="rate_limit_exceeded",
                    client_ip=client_ip,
                    user_id=resolved_user_id,
                    details=error or "Rate limit exceeded",
                )
                raise HTTPException(status_code=429, detail=error)

        # File validation if raw bytes are provided
        if file_bytes:
            content_type = content_type or (
                "application/pdf" if safe_filename.lower().endswith(".pdf") else "application/octet-stream"
            )
            is_valid, error = SecurityManager.validate_file(file_bytes, safe_filename, content_type)
            if not is_valid:
                logger.warning(f"[ParserService] File validation failed: {error}")
                SecurityAuditLogger.log_security_violation(
                    violation_type="invalid_file",
                    client_ip=client_ip,
                    user_id=resolved_user_id,
                    details=error or "File validation failed",
                )
                raise HTTPException(status_code=400, detail=error)

        logger.info(f"[ParserService] Starting resume extraction for user '{resolved_user_id}' (filename='{safe_filename}')...")

        # ── Execution Path 1: Direct Document Bytes (PDF, DOCX, TXT) ──
        if file_bytes:
            fn_lower = safe_filename.lower()

            if fn_lower.endswith(".pdf"):
                result = await ImprovedUniversalResumeParser.parse_document_bytes(
                    file_bytes=file_bytes,
                    content_type="application/pdf",
                    filename=safe_filename,
                    user_id=resolved_user_id,
                )
            elif fn_lower.endswith((".docx", ".doc")):
                try:
                    import docx
                    doc = docx.Document(io.BytesIO(file_bytes))
                    doc_text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                            if row_text:
                                doc_text_parts.append(row_text)
                    extracted_text = "\n".join(doc_text_parts)
                except Exception as docx_err:
                    logger.warning(f"[ParserService] python-docx parsing failed ({docx_err}), using text decoding...")
                    extracted_text = file_bytes.decode("utf-8", errors="ignore")

                result = await ImprovedUniversalResumeParser.parse_text(
                    text_content=extracted_text,
                    user_id=resolved_user_id,
                )
            else:
                # Plain text / other supported text formats
                raw_text = file_bytes.decode("utf-8", errors="ignore")
                result = await ImprovedUniversalResumeParser.parse_text(
                    text_content=raw_text,
                    user_id=resolved_user_id,
                )

        # ── Execution Path 2: Extractor Output / Text Payload ──
        elif extractor_output is not None:
            result = await ImprovedUniversalResumeParser.parse(extractor_output)
        else:
            raise HTTPException(status_code=400, detail="No resume content or file bytes provided for parsing")

        if not result.get("success", False):
            logger.error("[ParserService] Resume parsing returned unsuccessful status")
            return {
                "success": False,
                "message": "Failed to parse resume",
                "parsed": None,
            }

        parsed_data = result.get("parsed", {})

        # Sanitize output
        sanitized_result = InputSanitizer.sanitize_resume_data(parsed_data)

        # Log successful file upload
        if file_bytes:
            SecurityAuditLogger.log_file_upload(
                filename=safe_filename,
                file_size=len(file_bytes),
                client_ip=client_ip,
                user_id=resolved_user_id,
                status="success",
                details=f"Exp: {len(sanitized_result.get('experience', []))}, Edu: {len(sanitized_result.get('education', []))}",
            )

        logger.info(f"[ParserService] ✓ Resume parsing successful (source: {result.get('source', 'gemini')})")
        return {
            "success": True,
            "message": "Resume parsed successfully",
            "parsed": sanitized_result,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[ParserService] Resume parsing error: {str(e)}", exc_info=True)
        return {
            "success": False,
            "message": f"Resume parsing failed: {str(e)}",
            "parsed": None,
        }