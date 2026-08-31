"""
resume_parser_helper.py — General-purpose layout-aware resume parsing helpers.

Provides:
- PyMuPDF block extraction preserving columnar layout.
- Python-docx paragraph extraction.
- Universal international contact extraction (global locations, international phones, multi-platform links).
- Keyword-based section boundary detection.
"""

import re
from typing import Any, Dict, List, Optional
import fitz  # PyMuPDF
from docx import Document


def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF using blocks to preserve layout better than plain text."""
    doc = fitz.open(pdf_path)
    full_text = ""
    for page in doc:
        blocks = page.get_text("blocks")
        for block in blocks:
            text_block = block[4].strip()
            if text_block:
                full_text += text_block + "\n\n"
    doc.close()
    return full_text.strip()


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from Word document preserving paragraphs and tables."""
    doc = Document(docx_path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts).strip()


def extract_personal_info(text: str) -> Dict[str, str]:
    """
    Extracts name, contact details, global location, and developer profile links via regex.
    Domain-agnostic and supports international candidate profiles worldwide.
    """
    info = {
        "name": "",
        "title": "",
        "email": "",
        "phone": "",
        "location": "",
        "portfolio": "",
        "github": "",
        "linkedin": "",
    }

    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # 1. Name heuristic (first clean line without numbers, @, or URLs)
    for line in lines[:5]:
        if not any(char.isdigit() for char in line) and "@" not in line and "http" not in line.lower() and len(line.split()) <= 4:
            info["name"] = line
            break

    # 2. Title heuristic (second line or line after name)
    if len(lines) > 1:
        for line in lines[1:6]:
            if line != info["name"] and "@" not in line and "http" not in line and len(line.split()) <= 6:
                info["title"] = line
                break

    # 3. Email
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    if email_match:
        info["email"] = email_match.group(0)

    # 4. International Phone (E.164 and common local formats)
    phone_match = re.search(r'(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,5}\)?[\s.-]?)?\d{3,5}[\s.-]?\d{3,5}[\s.-]?\d{0,5}', text)
    if phone_match and len(re.findall(r'\d', phone_match.group(0))) >= 7:
        info["phone"] = phone_match.group(0).strip()

    # 5. Universal Location Matching (City, State/Country/Postal)
    loc_match = re.search(r'\b([A-Z][a-zA-Z\s]+,\s*[A-Z]{2}|[A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z\s]+(?:\s+\d{5})?)\b', text[:1000])
    if loc_match:
        info["location"] = loc_match.group(0).strip()

    # 6. Profile & Developer Links
    github = re.search(r'(https?://)?(www\.)?(?:github\.com|gitlab\.com)/[^\s,]+', text, re.IGNORECASE)
    if github:
        info["github"] = github.group(0)

    linkedin = re.search(r'(https?://)?(www\.)?linkedin\.com/in/[^\s,]+', text, re.IGNORECASE)
    if linkedin:
        info["linkedin"] = linkedin.group(0)

    portfolio = re.search(r'(https?://)?(?:www\.)?(?:[a-zA-Z0-9-]+\.(?:io|me|dev|app|org|com)/[^\s,]*|kaggle\.com/[^\s,]+|huggingface\.co/[^\s,]+)', text, re.IGNORECASE)
    if portfolio and not any(k in portfolio.group(0).lower() for k in ("github.com", "linkedin.com")):
        info["portfolio"] = portfolio.group(0)

    return info


def parse_resume_sections(text: str) -> List[Dict[str, Any]]:
    """
    Robust section detection using generalized keyword matching.
    """
    section_keywords = {
        "summary": ["summary", "professional summary", "profile summary", "about me", "objective"],
        "experience": ["professional experience", "work experience", "experience", "employment", "work history"],
        "education": ["education", "academic background", "academic qualifications", "academic", "qualifications"],
        "skills": ["technical skills", "skills", "technical stack", "tech stack", "core competencies", "technologies"],
        "projects": ["projects", "key projects", "personal projects", "academic projects"],
        "certifications": ["certifications", "certificates", "licenses"],
        "languages": ["languages", "language proficiency"],
        "publications": ["publications", "research"],
    }

    sections: List[Dict[str, Any]] = []
    current_section = {"heading": "HEADER", "content": []}
    lines = text.split("\n")

    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue

        lower_line = clean_line.lower()

        matched_section = None
        for sec_key, keywords in section_keywords.items():
            if any(lower_line.startswith(kw) or lower_line == kw for kw in keywords):
                matched_section = sec_key
                break

        if matched_section:
            if current_section["content"]:
                sections.append({
                    "heading": current_section["heading"],
                    "content": "\n".join(current_section["content"]).strip(),
                })
            current_section = {
                "heading": clean_line,
                "content": [],
            }
        else:
            current_section["content"].append(clean_line)

    if current_section["content"]:
        sections.append({
            "heading": current_section["heading"],
            "content": "\n".join(current_section["content"]).strip(),
        })

    return sections


def parsing_resume(file_path: str, file_ext: str) -> Dict[str, Any]:
    """Main parser helper returning structured JSON."""
    if file_ext.lower() == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif file_ext.lower() in (".docx", ".doc"):
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX allowed.")

    personal_info = extract_personal_info(text)
    sections = parse_resume_sections(text)

    return {
        "status": "success",
        "personal_info": personal_info,
        "sections": sections,
        "raw_text": text[:10000] + "..." if len(text) > 10000 else text,
        "raw_text_length": len(text),
    }